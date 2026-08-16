from __future__ import annotations

import math
import re
from collections import Counter
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from pypdf import PdfReader

from search import is_job_visible_now
from taxonomy_engine import Taxonomy, normalize, tokens

TZ = ZoneInfo("Africa/Casablanca")

MAX_CV_BYTES = 5 * 1024 * 1024
MAX_CV_PAGES = 25
MAX_TEXT_CHARS = 120_000
SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt"})

CV_STOPWORDS = frozenset({
    "curriculum", "vitae", "cv", "profil", "profile", "resume", "contact",
    "telephone", "phone", "email", "adresse", "address", "maroc", "morocco",
    "experience", "experiences", "professionnelle", "professionnelles",
    "professional", "formation", "formations", "education", "diplome",
    "diplomes", "diploma", "diplomas", "competence", "competences", "skills",
    "skill", "langue", "langues", "language", "languages", "mission", "missions",
    "poste", "postes", "travail", "work", "emploi", "job", "jobs", "societe",
    "company", "entreprise", "annee", "annees", "year", "years", "mois", "month",
    "months", "niveau", "nationalite", "nationality", "permis", "driving",
})

LANGUAGE_MARKERS = {
    "ar": ("arabe", "arabic", "العربية", "عربي"),
    "fr": ("francais", "français", "french", "الفرنسية", "فرنسي"),
    "en": ("anglais", "english", "الانجليزية", "الإنجليزية", "إنجليزي"),
    "es": ("espagnol", "spanish", "الإسبانية", "اسباني"),
    "de": ("allemand", "german", "الألمانية", "الالمانية"),
}

EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", re.I)
PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)")
YEARS_RE = re.compile(
    r"(?<!\d)(\d{1,2})\s*(?:ans?|années?|annees?|years?|yrs?|سنة|سنوات)(?!\w)",
    re.I,
)


# CV-only contextual equivalences. These do NOT modify Masari's source-job
# taxonomy/classifier. They translate common CV wording into existing canonical
# professions so a résumé can express a role differently from a job advert.

CV_DIPLOMA_ROLE_RULES = (
    (
        (
            "technicien specialise en gros oeuvres",
            "technicien spécialisé en gros œuvres",
            "technicien specialise gros oeuvres",
            "technicien spécialisé gros œuvres",
        ),
        (
            ("btp.technicien_batiment", 100.0),
            ("btp.technicien_genie_civil", 94.0),
        ),
    ),
    (
        (
            "technicien dessinateur batiment",
            "technicien dessinateur bâtiment",
            "technicien en dessin de batiment",
            "technicien en dessin de bâtiment",
        ),
        (
            ("btp.dessinateur_architectural", 100.0),
            ("btp.technicien_bureau_etudes", 92.0),
            ("btp.dessinateur_projeteur", 88.0),
        ),
    ),
)

CV_CONTEXT_ROLE_RULES = (
    (
        (
            "gestionnaire de chantier",
            "gestionnaire chantier",
            "gestion de chantier",
            "conduite de chantier",
            "suivi de chantier",
            "coordination de chantier",
            "coordination des travaux",
        ),
        (
            ("btp.conducteur_travaux", 100.0),
            ("btp.chef_chantier", 96.0),
            ("btp.technicien_genie_civil", 88.0),
        ),
    ),
    (
        ("gros oeuvre", "gros œuvres", "gros œuvre"),
        (
            ("btp.technicien_batiment", 92.0),
            ("btp.technicien_genie_civil", 90.0),
        ),
    ),
    (
        (
            "metre",
            "metre batiment",
            "metres",
            "métré",
            "métrés",
            "calcul de metre",
            "calcul de metres",
            "estimation des couts",
            "estimation des coûts",
        ),
        (("btp.metreur", 94.0),),
    ),
    (
        (
            "dessinateur batiment",
            "dessinateur bâtiment",
            "dessin batiment",
            "dessin bâtiment",
            "dessinateur architecture",
            "autocad",
            "archicad",
            "dao cao",
        ),
        (
            ("btp.dessinateur_architectural", 93.0),
            ("btp.dessinateur_projeteur", 84.0),
        ),
    ),
)

# Generic labels that are misleading when they occur inside a clearly
# technical role phrase.
CV_CONTEXT_BLOCKS = {
    "admin.gestionnaire": (
        "gestionnaire de chantier",
        "gestionnaire chantier",
    ),
    "admin.assistant": (
        "assistant technicien",
        "assistant technique",
    ),
}

CV_SECTION_HEADERS = {
    "profil professionnel": "profile",
    "profile": "profile",
    "experiences professionnelles": "experience",
    "expériences professionnelles": "experience",
    "professional experience": "experience",
    "formation": "education",
    "education": "education",
    "competences techniques": "skills",
    "compétences techniques": "skills",
    "technical skills": "skills",
    "langues": "languages",
    "languages": "languages",
}


class CVAnalysisError(ValueError):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def _safe_extension(filename: str) -> str:
    return Path(str(filename or "")).suffix.casefold()


def _clean_text(text: str) -> str:
    text = str(text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()[:MAX_TEXT_CHARS]


def _extract_pdf(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise CVAnalysisError("CV_PDF_INVALID", "Impossible de lire ce fichier PDF.") from exc

    if len(reader.pages) > MAX_CV_PAGES:
        raise CVAnalysisError(
            "CV_TOO_MANY_PAGES",
            f"Le CV dépasse la limite de {MAX_CV_PAGES} pages.",
        )

    if getattr(reader, "is_encrypted", False):
        try:
            result = reader.decrypt("")
        except Exception as exc:
            raise CVAnalysisError("CV_PDF_ENCRYPTED", "Le PDF est protégé par mot de passe.") from exc
        if result == 0:
            raise CVAnalysisError("CV_PDF_ENCRYPTED", "Le PDF est protégé par mot de passe.")

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")

    return _clean_text("\n".join(pages)), {"pages": len(reader.pages)}


def _extract_docx(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise CVAnalysisError("CV_DOCX_INVALID", "Impossible de lire ce fichier DOCX.") from exc

    blocks = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                blocks.append(" | ".join(values))

    return _clean_text("\n".join(blocks)), {"paragraphs": len(document.paragraphs)}


def _extract_txt(data: bytes) -> tuple[str, dict[str, Any]]:
    text = ""
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    return _clean_text(text), {}


def extract_cv_text(filename: str, data: bytes) -> tuple[str, dict[str, Any]]:
    if not isinstance(data, (bytes, bytearray)):
        raise CVAnalysisError("CV_INVALID_BYTES", "Fichier CV invalide.")
    if not data:
        raise CVAnalysisError("CV_EMPTY_FILE", "Le fichier CV est vide.")
    if len(data) > MAX_CV_BYTES:
        raise CVAnalysisError(
            "CV_TOO_LARGE",
            f"Le CV dépasse la limite de {MAX_CV_BYTES // (1024 * 1024)} Mo.",
        )

    ext = _safe_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise CVAnalysisError(
            "CV_UNSUPPORTED_FORMAT",
            "Formats acceptés : PDF, DOCX et TXT.",
        )

    if ext == ".pdf":
        text, meta = _extract_pdf(bytes(data))
    elif ext == ".docx":
        text, meta = _extract_docx(bytes(data))
    else:
        text, meta = _extract_txt(bytes(data))

    if len(normalize(text)) < 60:
        raise CVAnalysisError(
            "CV_TEXT_TOO_SHORT",
            "Le CV ne contient pas assez de texte exploitable. "
            "Pour un PDF scanné, utilisez une version contenant du texte.",
        )

    return text, {"extension": ext, "bytes": len(data), **meta}


def _analysis_text(text: str) -> str:
    text = EMAIL_RE.sub(" ", text)
    text = PHONE_RE.sub(" ", text)
    return text


def detect_language_mentions(text: str) -> list[str]:
    normalized = normalize(text)
    padded = f" {normalized} "
    found = []
    for code, markers in LANGUAGE_MARKERS.items():
        for marker in markers:
            n = normalize(marker)
            if n and f" {n} " in padded:
                found.append(code)
                break
    return found


def infer_explicit_experience_years(text: str) -> int | None:
    values = []
    for match in YEARS_RE.finditer(text):
        value = int(match.group(1))
        if 0 < value <= 50:
            values.append(value)
    return max(values) if values else None


def _cv_lines(text: str) -> list[tuple[str, str, float]]:
    section = "headline"
    rows: list[tuple[str, str, float]] = []
    seen_structured_header = False
    preheader_line_count = 0

    for raw in _analysis_text(text).splitlines():
        line = raw.strip(" •\t")
        if not line:
            continue
        nline = normalize(line)

        if nline in CV_SECTION_HEADERS:
            section = CV_SECTION_HEADERS[nline]
            seen_structured_header = True
            continue

        # Before the first real CV section, only the first two non-empty lines
        # are treated as headline/title evidence (typically name + target role).
        # Longer summary prose must not receive headline strength.
        if not seen_structured_header:
            section = "headline" if preheader_line_count < 2 else "body"
            preheader_line_count += 1

        if section == "headline":
            weight = 1.35
        elif section == "profile":
            weight = 1.22
        elif section == "experience":
            weight = 1.30
        elif section == "education":
            weight = 1.12
        elif section == "skills":
            weight = 1.00
        else:
            weight = 0.82

        current_markers = (
            "actuellement", "depuis", "en cours",
            "present", "présent", "current",
        )
        if section in {"headline", "profile", "experience"} and any(
            marker in nline for marker in current_markers
        ):
            weight += 0.28

        if section in {"experience", "education"} and len(tokens(line)) <= 10:
            weight += 0.12

        rows.append((line, section, weight))

    return rows


def _blocked_generic_hit(profession_id: str, line: str) -> bool:
    blocked_phrases = CV_CONTEXT_BLOCKS.get(profession_id) or ()
    nline = normalize(line)
    return any(normalize(phrase) in nline for phrase in blocked_phrases)


def _profession_candidates(
    text: str,
    taxonomy: Taxonomy,
    *,
    limit: int = 6,
) -> list[dict[str, Any]]:
    lines = _cv_lines(text)
    full_normalized = normalize(_analysis_text(text))
    full_padded = f" {full_normalized} "

    scores: dict[str, float] = {}
    evidence: dict[str, set[str]] = {}
    evidence_kinds: dict[str, set[str]] = {}

    def add(pid: str, score: float, value: str, kind: str) -> None:
        profession = taxonomy.profession(pid)
        if profession is None:
            return
        scores[pid] = max(scores.get(pid, 0.0), min(float(score), 100.0))
        if value:
            evidence.setdefault(pid, set()).add(value)
        evidence_kinds.setdefault(pid, set()).add(kind)

    # 1) Section-aware exact/strong profession vocabulary and diploma evidence.
    for line, section, weight in lines:
        nline = normalize(line)
        padded = f" {nline} "
        line_token_count = len(tokens(line))
        title_like = line_token_count <= 10

        for profession in taxonomy.market:
            if _blocked_generic_hit(profession.id, line):
                continue

            for term in taxonomy.query_terms(profession):
                nt = normalize(term)
                if not nt or f" {nt} " not in padded:
                    continue

                term_token_count = len(tokens(nt))
                # Single generic words are accepted only in title-like lines.
                if term_token_count <= 1 and not title_like:
                    continue

                base_score = 79.0 if term_token_count >= 2 else 68.0
                if section == "experience":
                    base_score += 7.0
                elif section == "headline":
                    base_score += 9.0
                elif section == "profile":
                    base_score += 4.0

                if normalize(term) == normalize(profession.label):
                    base_score += 4.0

                add(
                    profession.id,
                    base_score * min(weight, 1.18),
                    term,
                    f"{section}:term",
                )

            # Diploma evidence is an eligibility signal, so it is accepted
            # only from the explicit education/formation section. Mentioning a
            # diploma name in profile prose must not create public eligibility.
            if section == "education":
                for diploma in profession.diplomas:
                    ndiploma = normalize(diploma)
                    if ndiploma and f" {ndiploma} " in padded:
                        add(
                            profession.id,
                            92.0 * min(weight, 1.08),
                            diploma,
                            "education:diploma",
                        )

        # Existing conservative query resolver is useful for short role headings.
        if title_like and section in {"headline", "experience", "education"}:
            for suggestion in taxonomy.autocomplete(line, limit=4):
                score = float(suggestion.get("score") or 0.0)
                pid = str(suggestion.get("profession_id") or "")
                if not pid or score < 800 or _blocked_generic_hit(pid, line):
                    continue
                mapped = 86.0 + min(10.0, (score - 800.0) / 40.0)
                if section == "experience":
                    mapped += 4.0
                add(
                    pid,
                    mapped * min(weight, 1.10),
                    str(suggestion.get("matched_term") or suggestion.get("label") or ""),
                    f"{section}:resolver",
                )

    # 2) Explicit diploma phrases from the Formation/Education section only.
    # These create the diploma-backed evidence used by public eligibility.
    for phrases, mappings in CV_DIPLOMA_ROLE_RULES:
        matched_phrases: list[str] = []
        for line, section, _weight in lines:
            if section != "education":
                continue
            nline = normalize(line)
            for phrase in phrases:
                if normalize(phrase) in nline:
                    matched_phrases.append(phrase)
        if not matched_phrases:
            continue
        for pid, base_score in mappings:
            add(
                pid,
                base_score,
                matched_phrases[0],
                "education:diploma_context",
            )

    # 3) CV-context aliases are section-aware. A current role/headline may
    # create strong profession evidence; the same words inside generic prose
    # are deliberately weaker and cannot override an explicit title.
    for phrases, mappings in CV_CONTEXT_ROLE_RULES:
        best_hit: tuple[float, str] | None = None
        for line, section, _weight in lines:
            nline = normalize(line)
            matched_phrase = next(
                (
                    phrase
                    for phrase in phrases
                    if normalize(phrase) in nline
                ),
                None,
            )
            if matched_phrase is None:
                continue

            if section in {"headline", "experience"}:
                factor = 1.00
            elif section == "profile":
                factor = 0.94
            elif section == "body":
                factor = 0.82
            elif section == "skills":
                factor = 0.78
            else:
                factor = 0.72

            current_markers = (
                "actuellement", "depuis", "en cours",
                "present", "présent", "current",
            )
            if section in {"profile", "experience"} and any(
                marker in nline for marker in current_markers
            ):
                factor = 1.00

            candidate = (factor, matched_phrase)
            if best_hit is None or candidate[0] > best_hit[0]:
                best_hit = candidate

        if best_hit is None:
            continue

        factor, matched_phrase = best_hit

        best_section = "body"
        best_current = False
        best_factor = -1.0
        for line, section, _weight in lines:
            nline = normalize(line)
            if not any(normalize(phrase) in nline for phrase in phrases):
                continue

            if section in {"headline", "experience"}:
                section_factor = 1.00
            elif section == "profile":
                section_factor = 0.94
            elif section == "body":
                section_factor = 0.82
            elif section == "skills":
                section_factor = 0.78
            else:
                section_factor = 0.72

            current_markers = (
                "actuellement", "depuis", "en cours",
                "present", "présent", "current",
            )
            is_current = (
                section in {"profile", "experience"}
                and any(marker in nline for marker in current_markers)
            )
            if is_current:
                section_factor = 1.00

            candidate_rank = (
                section_factor,
                1 if is_current else 0,
                1 if section in {"headline", "experience"} else 0,
            )
            current_rank = (
                best_factor,
                1 if best_current else 0,
                1 if best_section in {"headline", "experience"} else 0,
            )
            if candidate_rank > current_rank:
                best_factor = section_factor
                best_section = section
                best_current = is_current

        context_kind = f"cv_context:{best_section}"
        if best_current:
            context_kind += ":current"

        for pid, base_score in mappings:
            add(
                pid,
                base_score * factor,
                matched_phrase,
                context_kind,
            )

    # 4) Repeated exact phrases can raise confidence slightly, but cannot create
    # an unrelated profession on their own.
    for pid in list(scores):
        profession = taxonomy.profession(pid)
        if profession is None:
            continue
        repeats = 0
        for term in taxonomy.query_terms(profession):
            nt = normalize(term)
            if nt and len(tokens(nt)) >= 2:
                repeats += full_padded.count(f" {nt} ")
        if repeats > 1:
            scores[pid] = min(100.0, scores[pid] + min(6.0, (repeats - 1) * 2.0))

    rows = []
    for pid, score in scores.items():
        if score < 70.0:
            continue
        profession = taxonomy.profession(pid)
        if profession is None:
            continue
        rows.append(
            {
                "profession_id": profession.id,
                "label": profession.label,
                "sector": profession.sector,
                "family": profession.family,
                "score": round(score, 1),
                "evidence": sorted(
                    evidence.get(pid, set()),
                    key=lambda x: (-len(x), x.casefold()),
                )[:6],
                "evidence_kinds": sorted(evidence_kinds.get(pid, set())),
            }
        )

    def evidence_authority(row: dict[str, Any]) -> int:
        kinds = {str(x) for x in (row.get("evidence_kinds") or [])}

        if any(
            kind.startswith("headline:term")
            or kind.startswith("headline:resolver")
            or kind.startswith("experience:term")
            or kind.startswith("experience:resolver")
            or kind.startswith("cv_context:headline")
            or kind.startswith("cv_context:experience:current")
            or kind.startswith("cv_context:profile:current")
            for kind in kinds
        ):
            return 5

        if any(kind.startswith("education:diploma") for kind in kinds):
            return 4

        if any(kind.startswith("cv_context:experience") for kind in kinds):
            return 4

        if any(
            kind.startswith("skills:")
            or kind.startswith("cv_context:skills")
            for kind in kinds
        ):
            return 3

        if any(
            kind.startswith("profile:")
            or kind.startswith("cv_context:profile")
            or kind.startswith("cv_context:body")
            for kind in kinds
        ):
            return 2

        return 1

    for row in rows:
        row["authority"] = evidence_authority(row)

    rows.sort(
        key=lambda x: (
            -int(x.get("authority") or 0),
            -float(x["score"]),
            x["label"].casefold(),
        )
    )

    rows = [row for row in rows if float(row["score"]) >= 70.0]

    selected = rows[:limit]
    selected_ids = {str(row["profession_id"]) for row in selected}

    # Public matching depends on diploma evidence. Never discard a
    # diploma-backed profession merely because current-role candidates filled
    # the display-oriented top-N.
    for row in rows:
        kinds = {str(x) for x in (row.get("evidence_kinds") or [])}
        diploma_backed = any(
            kind.startswith("education:diploma")
            for kind in kinds
        )
        pid = str(row["profession_id"])
        if diploma_backed and pid not in selected_ids:
            selected.append(row)
            selected_ids.add(pid)

    selected.sort(
        key=lambda x: (
            -int(x.get("authority") or 0),
            -float(x["score"]),
            x["label"].casefold(),
        )
    )
    return selected

def _content_tokens(value: str) -> list[str]:
    result = []
    for token in tokens(value):
        if token in CV_STOPWORDS:
            continue
        if token.isdigit():
            continue
        if len(token) < 3 and token not in {"c++", "c#", "r"}:
            continue
        result.append(token)
    return result


def _flatten(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, dict):
        return " ".join(_flatten(v) for v in value.values())
    if isinstance(value, (list, tuple, set)):
        return " ".join(_flatten(v) for v in value)
    return str(value)


def _job_text(job: dict[str, Any]) -> str:
    fields = (
        "title",
        "job_name",
        "listing_title",
        "search_title",
        "specialties",
        "description",
        "profile",
        "education",
        "experience",
        "languages",
        "sector",
        "contract_type",
    )
    return " ".join(_flatten(job.get(field)) for field in fields)


def _job_profession_ids(job: dict[str, Any]) -> set[str]:
    ids = {
        str(x).strip()
        for x in (job.get("profession_ids") or [])
        if str(x).strip()
    }
    for match in job.get("profession_matches") or []:
        pid = str((match or {}).get("profession_id") or "").strip()
        if not pid:
            continue
        if (match or {}).get("searchable") is False:
            continue
        if (match or {}).get("confidence") == "RELATED":
            continue
        ids.add(pid)
    return ids


def _publication_age_days(job: dict[str, Any], now: datetime) -> int:
    raw = str(job.get("publication_date") or "")
    if not raw:
        return 99
    try:
        dt = datetime.fromisoformat(raw)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=TZ)
        else:
            dt = dt.astimezone(TZ)
        return max((now.date() - dt.date()).days, 0)
    except Exception:
        return 99


def _public_job_match(job: dict[str, Any], score: float, reasons: list[dict[str, Any]]) -> dict[str, Any]:
    gid = str(job.get("global_id") or job.get("uuid") or "")
    return {
        "global_id": gid,
        "title": str(job.get("job_name") or job.get("title") or job.get("search_title") or "").strip(),
        "company": str(job.get("company") or job.get("administration") or "").strip(),
        "source": str(job.get("source") or "").strip(),
        "source_label": str(job.get("source_label") or job.get("source") or "").strip(),
        "employment_sector": str(job.get("employment_sector") or job.get("scope") or "").strip(),
        "location": _flatten(job.get("location") or job.get("work_location_text")).strip(),
        "publication_date": job.get("publication_date"),
        "deadline": job.get("deadline"),
        "contract_type": job.get("contract_type") or job.get("recruitment_type"),
        "match_score": round(score, 1),
        "reasons": reasons,
        "detail_url": f"/job/{gid}",
    }


def _profession_evidence_flags(
    profession_rows: list[dict[str, Any]],
) -> dict[str, dict[str, bool]]:
    out: dict[str, dict[str, bool]] = {}
    for row in profession_rows:
        pid = str(row.get("profession_id") or "")
        kinds = {str(x) for x in (row.get("evidence_kinds") or [])}
        out[pid] = {
            "diploma": any(
                kind.startswith("education:diploma")
                for kind in kinds
            ),
            "experience": any(
                kind.startswith("experience:")
                or kind.startswith("cv_context:experience")
                or kind.startswith("cv_context:profile:current")
                for kind in kinds
            ),
            "headline": any(
                kind.startswith("headline:")
                or kind.startswith("cv_context:headline")
                for kind in kinds
            ),
        }
    return out


def _relation_evidence(
    relation_pid: str,
    profession_rows: list[dict[str, Any]],
    taxonomy: Taxonomy,
) -> dict[str, bool]:
    flags = _profession_evidence_flags(profession_rows)
    direct = flags.get(relation_pid) or {}
    if direct.get("diploma") or direct.get("experience") or direct.get("headline"):
        return {
            "diploma": bool(direct.get("diploma")),
            "experience": bool(direct.get("experience")),
            "headline": bool(direct.get("headline")),
        }

    # Related canonical professions can inherit evidence only when the HCP
    # family/root is genuinely compatible.
    target = taxonomy.profession(relation_pid)
    if target is None:
        return {"diploma": False, "experience": False, "headline": False}

    inherited = {"diploma": False, "experience": False, "headline": False}
    for pid, evidence in flags.items():
        prof = taxonomy.profession(pid)
        if prof is None:
            continue
        same_root = _sector_root(prof.sector) == _sector_root(target.sector)
        shared_hcp = bool(set(prof.hcp_codes) & set(target.hcp_codes))
        if same_root and shared_hcp:
            for key in inherited:
                inherited[key] = inherited[key] or bool(evidence.get(key))
    return inherited


def _sector_for_job(job: dict[str, Any]) -> str:
    explicit = str(job.get("employment_sector") or job.get("scope") or "").strip().casefold()
    if explicit in {"public", "private"}:
        return explicit
    source = str(job.get("source") or "").strip().casefold()
    if source == "emploi-public":
        return "public"
    return "private"


def _sector_root(value: str) -> str:
    return normalize(str(value or "").split("/", 1)[0])


def _profession_relation(
    job_pids: set[str],
    profession_rows: list[dict[str, Any]],
    taxonomy: Taxonomy,
) -> tuple[str, str, float] | None:
    if not job_pids or not profession_rows:
        return None

    best: tuple[str, str, float] | None = None
    best_rank: tuple[int, float] = (-1, -1.0)

    for cv_row in profession_rows:
        cv_pid = str(cv_row.get("profession_id") or "")
        cv_score = float(cv_row.get("score") or 0.0)
        cv_prof = taxonomy.profession(cv_pid)
        if cv_prof is None:
            continue

        for job_pid in job_pids:
            job_prof = taxonomy.profession(job_pid)
            if job_prof is None:
                continue

            if job_pid == cv_pid:
                candidate = ("exact", cv_pid, cv_score)
            else:
                same_root = (
                    _sector_root(job_prof.sector)
                    and _sector_root(job_prof.sector) == _sector_root(cv_prof.sector)
                )
                shared_hcp = bool(set(job_prof.hcp_codes) & set(cv_prof.hcp_codes))
                same_family = bool(
                    normalize(job_prof.family)
                    and normalize(job_prof.family) == normalize(cv_prof.family)
                )

                if same_root and same_family:
                    candidate = ("related", cv_pid, cv_score * 0.86)
                elif same_root and shared_hcp:
                    candidate = ("related", cv_pid, cv_score * 0.78)
                else:
                    continue

            # Exact profession identity is authoritative over a merely
            # related profession even when the related CV role has a slightly
            # higher profile score. This is essential for diploma eligibility.
            priority = 2 if candidate[0] == "exact" else 1
            rank = (priority, candidate[2])
            if best is None or rank > best_rank:
                best = candidate
                best_rank = rank

    return best


def _match_jobs(
    text: str,
    index: dict[str, Any],
    profession_rows: list[dict[str, Any]],
    taxonomy: Taxonomy,
    *,
    limit: int,
    now: datetime,
) -> tuple[
    list[dict[str, Any]],
    list[str],
    dict[str, list[dict[str, Any]]],
]:
    visible_jobs = [
        job
        for job in (index.get("jobs") or [])
        if isinstance(job, dict) and is_job_visible_now(job, now)
    ]
    if not visible_jobs or not profession_rows:
        return [], [], {"public": [], "private": []}

    cv_counts = Counter(_content_tokens(_analysis_text(text)))
    if not cv_counts:
        return [], [], {"public": [], "private": []}

    job_token_sets: list[tuple[dict[str, Any], set[str]]] = []
    df: Counter[str] = Counter()
    cv_token_set = set(cv_counts)

    for job in visible_jobs:
        job_tokens = set(_content_tokens(_job_text(job)))
        job_token_sets.append((job, job_tokens))
        for token in job_tokens & cv_token_set:
            df[token] += 1

    n_jobs = max(len(visible_jobs), 1)
    idf = {
        token: math.log((n_jobs + 1) / (count + 1)) + 1.0
        for token, count in df.items()
        if 0 < count / n_jobs <= 0.22
    }
    cv_weights = {
        token: min(float(cv_counts[token]), 3.0) * weight
        for token, weight in idf.items()
    }
    strongest_cv_tokens = dict(
        sorted(cv_weights.items(), key=lambda x: (-x[1], x[0]))[:70]
    )
    denominator = sum(strongest_cv_tokens.values()) or 1.0

    labels = {
        str(row["profession_id"]): str(row["label"])
        for row in profession_rows
    }

    scored: list[
        tuple[float, str, str, dict[str, Any], dict[str, Any]]
    ] = []

    for job, job_tokens in job_token_sets:
        job_pids = _job_profession_ids(job)
        relation = _profession_relation(job_pids, profession_rows, taxonomy)

        # Precision gate: keyword-only matches are not job recommendations.
        if relation is None:
            continue

        relation_kind, cv_pid, relation_score = relation
        matched = [token for token in strongest_cv_tokens if token in job_tokens]
        sector = _sector_for_job(job)
        evidence = _relation_evidence(cv_pid, profession_rows, taxonomy)

        # Morocco policy model:
        # - Public: diploma-backed professional eligibility is mandatory.
        #   Experience/keywords can rank an already eligible candidate, but
        #   cannot make an ineligible diploma profile eligible.
        # - Private: profession relation is mandatory, then experience,
        #   diploma and technical overlap jointly rank the result.
        if sector == "public" and not evidence["diploma"]:
            continue

        if relation_kind == "related" and not matched:
            continue

        matched_weight = sum(strongest_cv_tokens[t] for t in matched)
        coverage = min(matched_weight / denominator, 1.0)

        title_tokens = set(
            _content_tokens(str(job.get("job_name") or job.get("title") or ""))
        )
        title_hits = title_tokens & set(matched)
        age = _publication_age_days(job, now)

        if sector == "public":
            diploma_component = 58.0 if relation_kind == "exact" else 48.0
            profession_component = min(
                22.0,
                relation_score * (0.22 if relation_kind == "exact" else 0.18),
            )
            keyword_component = min(
                8.0,
                (min(len(matched), 3) * 1.5) + (coverage * 4.0),
            )
            title_component = min(7.0, len(title_hits) * 2.0)
            experience_component = 3.0 if evidence["experience"] else 0.0
            recency_component = max(
                0.0,
                2.0 * (15 - min(age, 15)) / 15.0,
            )
            score = min(
                100.0,
                diploma_component
                + profession_component
                + keyword_component
                + title_component
                + experience_component
                + recency_component,
            )
            minimum_score = 68.0
        else:
            profession_component = min(
                46.0,
                relation_score * (0.46 if relation_kind == "exact" else 0.38),
            )
            experience_component = 20.0 if evidence["experience"] else (
                8.0 if evidence["headline"] else 0.0
            )
            diploma_component = 12.0 if evidence["diploma"] else 0.0
            keyword_component = min(
                16.0,
                (min(len(matched), 5) * 1.8) + (coverage * 7.0),
            )
            title_component = min(4.0, len(title_hits) * 1.5)
            recency_component = max(
                0.0,
                2.0 * (15 - min(age, 15)) / 15.0,
            )
            score = min(
                100.0,
                profession_component
                + experience_component
                + diploma_component
                + keyword_component
                + title_component
                + recency_component,
            )
            minimum_score = 48.0

        if score < minimum_score:
            continue

        reasons: list[dict[str, Any]] = [
            {
                "type": (
                    "profession"
                    if relation_kind == "exact"
                    else "related_profession"
                ),
                "label": labels.get(cv_pid, cv_pid),
            }
        ]
        if sector == "public":
            reasons.append({"type": "diploma_eligibility"})
        elif evidence["experience"]:
            reasons.append({"type": "experience_relevance"})
        if sector == "private" and evidence["diploma"]:
            reasons.append({"type": "diploma_support"})

        if matched:
            matched_sorted = sorted(
                matched,
                key=lambda t: (-strongest_cv_tokens.get(t, 0.0), t),
            )[:6]
            reasons.append({"type": "keywords", "terms": matched_sorted})

        if age <= 3:
            reasons.append({"type": "recent", "days": age})

        public = _public_job_match(job, score, reasons)
        public["employment_sector"] = sector

        scored.append(
            (
                score,
                str(job.get("publication_date") or ""),
                str(job.get("global_id") or job.get("uuid") or ""),
                public,
                job,
            )
        )

    scored.sort(key=lambda x: x[1], reverse=True)
    scored.sort(key=lambda x: x[0], reverse=True)

    sector_matches = {"public": [], "private": []}
    for row in scored:
        sector = str(row[3].get("employment_sector") or "")
        if sector not in sector_matches:
            sector = "private"
        if len(sector_matches[sector]) < limit:
            sector_matches[sector].append(row[3])

    # Backward-compatible global top-N. The UI v2 uses sector_matches.
    matches = [row[3] for row in scored[:limit]]

    cv_tokens_all = set(_content_tokens(_analysis_text(text)))
    review_counts: Counter[str] = Counter()
    for _score, _date, _gid, _public, job in scored[:12]:
        requirement_text = " ".join(
            _flatten(job.get(field))
            for field in ("specialties", "profile", "education", "experience")
        )
        for token in set(_content_tokens(requirement_text)):
            if token not in cv_tokens_all and len(token) >= 4:
                review_counts[token] += 1

    review_terms = [
        token
        for token, count in review_counts.most_common(20)
        if count >= 2
    ][:8]

    return matches, review_terms, sector_matches

def analyze_cv(
    filename: str,
    data: bytes,
    index: dict[str, Any],
    taxonomy: Taxonomy,
    *,
    limit: int = 15,
    now: datetime | None = None,
) -> dict[str, Any]:
    limit = max(1, min(int(limit), 15))
    now = now or datetime.now(TZ)
    if now.tzinfo is None:
        now = now.replace(tzinfo=TZ)
    else:
        now = now.astimezone(TZ)

    text, file_meta = extract_cv_text(filename, data)
    safe_text = _analysis_text(text)

    # Matching needs a wider candidate set than the compact profile display.
    # In particular, diploma-backed secondary professions must remain available
    # for public-sector eligibility even when the current role is ranked first.
    profession_candidates = _profession_candidates(
        safe_text,
        taxonomy,
        limit=12,
    )
    profile_professions = profession_candidates[:8]

    matches, review_terms, sector_matches = _match_jobs(
        safe_text,
        index,
        profession_candidates,
        taxonomy,
        limit=limit,
        now=now,
    )

    keyword_counts = Counter(_content_tokens(safe_text))
    keywords = [
        token
        for token, _count in keyword_counts.most_common(30)
        if len(token) >= 3
    ][:15]

    top_prof_score = (
        float(profile_professions[0]["score"])
        if profile_professions
        else 0.0
    )
    if len(safe_text) >= 700 and top_prof_score >= 88:
        confidence = "high"
    elif len(safe_text) >= 300 and top_prof_score >= 75:
        confidence = "medium"
    else:
        confidence = "low"

    warnings = []
    if len(safe_text) < 300:
        warnings.append("CV_TEXT_SHORT")
    if not profile_professions:
        warnings.append("NO_STRONG_PROFESSION_DETECTED")
    if not sector_matches["public"] and not sector_matches["private"]:
        warnings.append("NO_JOB_MATCHES")

    all_sector_matches = [
        *sector_matches["public"],
        *sector_matches["private"],
    ]
    source_counts = Counter(match["source"] for match in all_sector_matches)

    return {
        "version": "cv-matching-v2",
        "file": {
            "name": Path(filename or "cv").name,
            **file_meta,
        },
        "profile": {
            "analysis_confidence": confidence,
            "professions": profile_professions,
            "keywords": keywords,
            "language_mentions": detect_language_mentions(safe_text),
            "explicit_experience_years": infer_explicit_experience_years(safe_text),
        },
        "matches": matches,
        "sector_matches": sector_matches,
        "sector_counts": {
            "public": len(sector_matches["public"]),
            "private": len(sector_matches["private"]),
        },
        "match_count": len(all_sector_matches),
        "source_counts": dict(sorted(source_counts.items())),
        "matching_policy": {
            "public": {
                "diploma_gate": True,
                "experience_can_override_diploma_gate": False,
                "model": "diploma_first",
            },
            "private": {
                "diploma_gate": False,
                "profession_gate": True,
                "model": "profession_experience_skills",
            },
        },
        "suggestions": {
            "terms_to_review": review_terms,
            "note": (
                "Ces termes apparaissent dans plusieurs offres correspondantes "
                "mais pas dans le texte du CV. Ils ne signifient pas que la compétence manque."
            ),
        },
        "warnings": warnings,
        "privacy": {
            "stored": False,
            "raw_text_returned": False,
            "processing": "memory_only",
        },
    }
